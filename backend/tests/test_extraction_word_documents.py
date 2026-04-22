"""
Regression tests for Word-document support in the extraction pipeline.

Background: before this change, `app.services.extraction.extract_text` dispatched
only spreadsheet / PDF / image. Any .docx attachment hit the fall-through branch
and returned `Unsupported MIME type`, which surfaced to the UI as
`no_structured_timesheet_data` — a misleading skip reason because the real
problem was that the extractor never attempted to read the file.

These tests exercise the new .docx path end-to-end (build a real .docx in
memory, hand it to `extract_text`, assert the text comes back) and assert
that the .doc path is at least wired up correctly (we don't invoke antiword
itself here — that's a system binary tested at runtime).
"""
from __future__ import annotations

import io

import pytest

from app.services.extraction import (
    WORD_DOC_MIME_TYPES,
    WORD_DOCX_MIME_TYPES,
    extract_text,
)


def _build_docx_bytes(
    paragraphs: list[str] | None = None,
    table_rows: list[list[str]] | None = None,
) -> bytes:
    """Build a minimal .docx in memory with the given paragraphs and table."""
    import docx

    document = docx.Document()
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_idx, row in enumerate(table_rows):
            for c_idx, value in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_extract_text_handles_docx_paragraphs():
    content = _build_docx_bytes(
        paragraphs=[
            "Timesheet for Daniel Gwilt",
            "Week of April 20, 2026",
            "Total hours: 40",
        ]
    )

    result = await extract_text(
        content=content,
        filename="daniel_gwilt_timesheet.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.method != "failed", f"extraction failed: {result.error}"
    assert "Daniel Gwilt" in result.text
    assert "Week of April 20, 2026" in result.text
    assert "Total hours: 40" in result.text


@pytest.mark.asyncio
async def test_extract_text_preserves_docx_table_row_structure():
    """
    Timesheets in Word are almost always tables. The extractor must emit row
    structure the downstream LLM can parse — cells joined by tabs, rows by
    newlines.
    """
    content = _build_docx_bytes(
        paragraphs=["Daniel Gwilt — April 20, 2026"],
        table_rows=[
            ["Date", "Hours", "Project"],
            ["2026-04-20", "8", "Acme Corp"],
            ["2026-04-21", "7.5", "Acme Corp"],
        ],
    )

    result = await extract_text(
        content=content,
        filename="daniel_gwilt_timesheet.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.method != "failed", f"extraction failed: {result.error}"
    # Table rows must land as distinct lines so the LLM sees per-row grouping.
    assert "2026-04-20\t8\tAcme Corp" in result.text
    assert "2026-04-21\t7.5\tAcme Corp" in result.text
    # Header row preserved.
    assert "Date\tHours\tProject" in result.text


@pytest.mark.asyncio
async def test_extract_text_infers_docx_from_suffix_when_mime_is_octet_stream():
    """
    Some mail clients deliver attachments with mime=application/octet-stream.
    The MIME normaliser falls back on the filename suffix — assert that path
    still routes .docx to the docx extractor.
    """
    content = _build_docx_bytes(paragraphs=["Fallback routing test"])

    result = await extract_text(
        content=content,
        filename="timesheet.docx",
        mime_type="application/octet-stream",
    )

    assert result.method != "failed", f"extraction failed: {result.error}"
    assert "Fallback routing test" in result.text


@pytest.mark.asyncio
async def test_extract_text_reports_failure_for_corrupt_docx():
    """A non-zip payload with a .docx name should fail cleanly, not crash."""
    result = await extract_text(
        content=b"this is not a real docx file",
        filename="broken.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.method == "failed"
    assert result.error  # non-empty error message


@pytest.mark.asyncio
async def test_extract_text_dedupes_horizontally_merged_cells():
    """
    python-docx reports a horizontally merged cell once per underlying grid
    column, each copy returning the full merged text. Before dedup, a 7-column
    row merged into a single cell produced `"Total Weekly Hours:" × 7` on one
    line. Assert the merged cell contributes its text exactly once.
    """
    import docx

    document = docx.Document()
    table = document.add_table(rows=2, cols=4)
    # Row 0: regular header row
    table.rows[0].cells[0].text = "Col A"
    table.rows[0].cells[1].text = "Col B"
    table.rows[0].cells[2].text = "Col C"
    table.rows[0].cells[3].text = "Col D"
    # Row 1: merge all four cells into one and set a single value.
    merged = table.rows[1].cells[0].merge(table.rows[1].cells[3])
    merged.text = "Total Weekly Hours: 40"

    buffer = io.BytesIO()
    document.save(buffer)

    result = await extract_text(
        content=buffer.getvalue(),
        filename="merged.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.method != "failed", f"extraction failed: {result.error}"
    # The merged-cell row must appear exactly once, not 4×.
    assert result.text.count("Total Weekly Hours: 40") == 1, (
        f"merged cell was duplicated in text output:\n{result.text!r}"
    )


@pytest.mark.asyncio
async def test_extract_text_emits_rendered_html_for_docx_tables():
    """
    The review panel already knows how to display `rendered_html` for
    spreadsheet attachments. For Word, we emit one <table> per docx-table so
    reviewers see a real table instead of a tab-separated wall of text.
    """
    content = _build_docx_bytes(
        paragraphs=["Daniel Gwilt — April 20, 2026"],
        table_rows=[
            ["Date", "Hours", "Project"],
            ["2026-04-20", "8", "Acme Corp"],
        ],
    )

    result = await extract_text(
        content=content,
        filename="timesheet.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.method != "failed"
    assert result.rendered_html is not None
    # Paragraph rendered, table rendered with the expected cell content.
    assert "<p>Daniel Gwilt" in result.rendered_html
    assert "<table" in result.rendered_html
    assert "<td>Date</td>" in result.rendered_html
    assert "<td>2026-04-20</td>" in result.rendered_html
    assert "<td>Acme Corp</td>" in result.rendered_html


@pytest.mark.asyncio
async def test_extract_text_emits_standalone_docx_html_document():
    """
    The review panel renders `rendered_html` inside a sandboxed <iframe>, which
    doesn't inherit the parent page's CSS. A bare <div>…</div> fragment renders
    in the browser's default color and looks unreadable on the dark theme.
    Guard against regressing to a fragment by asserting the output is a full
    standalone document with styling — same approach xlsx_render uses.
    """
    content = _build_docx_bytes(paragraphs=["Hello"])

    result = await extract_text(
        content=content,
        filename="styled.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.rendered_html is not None
    html_lower = result.rendered_html.lower()
    assert "<!doctype html>" in html_lower
    assert "<style>" in html_lower
    # Dark-mode rule is what actually fixed the user-visible bug; assert it.
    assert "prefers-color-scheme: dark" in result.rendered_html


@pytest.mark.asyncio
async def test_extract_text_escapes_html_in_docx_content():
    """
    Whatever the user types into a Word document must be HTML-escaped on the
    way into `rendered_html`. A user typing `<script>alert(1)</script>` into a
    cell must not produce a live script tag in the review panel.
    """
    content = _build_docx_bytes(
        paragraphs=["Notes: <script>alert('xss')</script>"],
        table_rows=[["<b>Bold?</b>", "Plain"]],
    )

    result = await extract_text(
        content=content,
        filename="xss.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result.method != "failed"
    assert result.rendered_html is not None
    assert "<script>" not in result.rendered_html
    assert "&lt;script&gt;" in result.rendered_html
    assert "<td>&lt;b&gt;Bold?&lt;/b&gt;</td>" in result.rendered_html


def test_word_mime_type_sets_do_not_overlap_with_existing_dispatch():
    """
    Sanity check: make sure we didn't accidentally claim a MIME type already
    owned by the spreadsheet or image branches.
    """
    from app.services.extraction import IMAGE_MIME_TYPES, SPREADSHEET_MIME_TYPES

    word_all = WORD_DOCX_MIME_TYPES | WORD_DOC_MIME_TYPES
    assert word_all.isdisjoint(SPREADSHEET_MIME_TYPES)
    assert word_all.isdisjoint(IMAGE_MIME_TYPES)
