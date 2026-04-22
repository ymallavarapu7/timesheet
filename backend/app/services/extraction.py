"""
Document extraction pipeline.
Extracts raw text (or structured Vision JSON) from PDF, spreadsheet, and image attachments.

Extraction order (port of extraction_service.ts):
  Spreadsheet  → native XLSX/CSV parser → text
  PDF          → native text (pdfplumber) → if > 100 chars done
                 else pdftoppm → Vision API (all pages) → JSON
                 else Tesseract OCR → text
  Image        → Vision API → JSON
                 else Tesseract OCR → text
"""

import csv
import io
import json
import logging
import re
from pathlib import Path
from typing import Literal

from app.core.config import settings

logger = logging.getLogger(__name__)


def _detect_tesseract_lang(text_sample: str) -> str:
    """Detect language from a text sample for Tesseract. Falls back to 'eng'."""
    try:
        from langdetect import detect
        lang_code = detect(text_sample)
        # Map common langdetect codes to Tesseract language codes
        LANG_MAP = {
            "en": "eng", "fr": "fra", "de": "deu", "es": "spa", "it": "ita",
            "pt": "por", "nl": "nld", "pl": "pol", "ru": "rus", "ja": "jpn",
            "zh-cn": "chi_sim", "zh-tw": "chi_tra", "ko": "kor", "ar": "ara",
            "hi": "hin", "tr": "tur", "vi": "vie", "th": "tha",
        }
        return LANG_MAP.get(lang_code, "eng")
    except Exception:
        return "eng"

ExtractionMethodType = Literal[
    "native_spreadsheet",
    "native_pdf",
    "tesseract",
    "vision_api",
    "failed",
]

OCR_CONFIDENCE_THRESHOLD = 70.0

SPREADSHEET_MIME_TYPES = {
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/csv",
    "application/csv",
}

IMAGE_MIME_TYPES = {
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/tiff",
    "image/bmp",
    "image/gif",
}

WORD_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# Legacy Word 97–2003 binary format. Can't be read by python-docx; handled via
# antiword subprocess when the system binary is available.
WORD_DOC_MIME_TYPES = {
    "application/msword",
    "application/vnd.ms-word",
    "application/doc",
    "application/x-msword",
}

# Vision prompt — copied exactly from extraction_service.ts
def _build_vision_prompt(reference_date: str | None = None) -> str:
    date_rule = ""
    if reference_date:
        date_rule = (
            f"For context: this document was received around {reference_date}. "
            "When the document shows a month and day but no explicit year, "
            "assume the year is the one that makes the timesheet period fall "
            "within ~90 days of the reference date — never default to an older "
            "year. Prefer the current or most recent matching year. "
        )
    return (
        "Extract all timesheet data from this image. "
        "The image may contain one timesheet or multiple weekly/monthly timesheets. "
        + date_rule +
        "For calendar-style weekly or monthly timesheets, treat the columns as "
        "Monday through Sunday when the header shows Mo/Tu/We/Th/Fr/Sa/Su. "
        "If the sheet shows only five working days with hours, map them to Monday "
        "through Friday unless weekend hours are explicitly entered. "
        "Do not create weekend line items when Saturday/Sunday cells show 0 or blank, "
        "and do not move Friday hours onto Saturday or Sunday. "
        "If the image is a summary or pivot sheet showing categories and totals "
        "rather than daily dated entries, do not invent a repeated work_date for "
        "each category row; instead infer the month/period from tokens like 2/26 "
        "and return an empty line_items array. "
        "Return only valid JSON with a top-level field timesheets, where timesheets "
        "is an array of objects with fields: employee_name, client_name, "
        "period_start (YYYY-MM-DD), period_end (YYYY-MM-DD), total_hours (number), "
        "line_items (array of {work_date, hours, description, project_code}), "
        "extraction_confidence (0-1), uncertain_fields (array of field names)."
    )


VISION_PROMPT = _build_vision_prompt()


def _validate_vision_timesheet(ts: dict) -> bool:
    """Basic structural validation of a vision-extracted timesheet."""
    if not isinstance(ts, dict):
        return False
    line_items = ts.get("line_items")
    if line_items is not None:
        if not isinstance(line_items, list):
            return False
        for item in line_items:
            if not isinstance(item, dict):
                return False
    return True


class ExtractionResult:
    def __init__(
        self,
        text: str,
        method: ExtractionMethodType,
        confidence: float | None = None,
        error: str | None = None,
        vision_timesheets: list[dict] | None = None,
        spreadsheet_preview: dict | None = None,
        rendered_html: str | None = None,
    ) -> None:
        self.text = text
        self.method = method
        self.confidence = confidence
        self.error = error
        # Populated when Vision API returns structured JSON directly.
        # When set, the pipeline skips a second LLM extraction call.
        self.vision_timesheets = vision_timesheets
        # Structured preview of spreadsheet content for the reviewer UI.
        # Shape: {"sheets": [{"name": str, "rows": list[list[str]]}]}.
        # Retained for backward-compat; new previews use `rendered_html`.
        self.spreadsheet_preview = spreadsheet_preview
        # Self-contained HTML document rendering the source spreadsheet for
        # reviewer display. Populated for xlsx/csv attachments.
        self.rendered_html = rendered_html

    @property
    def success(self) -> bool:
        return self.method != "failed" and (
            bool(self.text.strip()) or bool(self.vision_timesheets)
        )


# ─── Spreadsheet extraction ───────────────────────────────────────────────────


def _trim_trailing_empty_columns(rows: list[list[str]]) -> list[list[str]]:
    """Drop columns from the right edge where every cell across all rows is empty."""
    if not rows:
        return rows
    width = max((len(r) for r in rows), default=0)
    keep = width
    while keep > 0 and all((row[keep - 1] if keep - 1 < len(row) else "").strip() == "" for row in rows):
        keep -= 1
    if keep == width:
        return rows
    return [row[:keep] for row in rows]


def _split_into_blocks(rows: list[list[str]]) -> list[list[list[str]]]:
    """Split rows horizontally on fully-empty separator columns.

    Excel sheets often pack multiple unrelated tables side-by-side with a blank
    column between them. This returns one rows-list per dense block.
    """
    if not rows:
        return []
    width = max((len(r) for r in rows), default=0)
    if width == 0:
        return []
    # A column is "empty" when no row has any non-whitespace content there.
    empty_cols = {
        c for c in range(width)
        if all(((row[c] if c < len(row) else "")).strip() == "" for row in rows)
    }
    # Walk columns left→right, accumulating dense ranges separated by empties.
    ranges: list[tuple[int, int]] = []
    start: int | None = None
    for c in range(width):
        if c in empty_cols:
            if start is not None:
                ranges.append((start, c))
                start = None
        else:
            if start is None:
                start = c
    if start is not None:
        ranges.append((start, width))
    blocks: list[list[list[str]]] = []
    for s, e in ranges:
        block = [row[s:e] for row in rows]
        # Drop rows that are fully empty inside this block.
        block = [row for row in block if any(cell.strip() for cell in row)]
        if block:
            blocks.append(block)
    return blocks


def _extract_spreadsheet(content: bytes, mime_type: str) -> ExtractionResult:
    try:
        if "csv" in mime_type:
            # Try to detect encoding, fallback to utf-8
            encoding = "utf-8"
            try:
                import chardet
                detected = chardet.detect(content[:10000])
                if detected and detected.get("encoding") and detected.get("confidence", 0) > 0.5:
                    encoding = detected["encoding"]
            except ImportError:
                pass
            text = content.decode(encoding, errors="replace")
            rows = [row for row in csv.reader(io.StringIO(text))]
            trimmed = _trim_trailing_empty_columns(rows)
            blocks = [{"rows": b} for b in _split_into_blocks(trimmed)]
            preview = {"sheets": [{"name": "Sheet1", "rows": trimmed, "blocks": blocks}]}
            csv_html = None
            try:
                from app.services.xlsx_render import render_csv_to_html
                csv_html = render_csv_to_html(content, encoding)
            except Exception as render_exc:
                logger.warning("csv HTML render failed (non-fatal): %s", render_exc)
            return ExtractionResult(
                text="\n".join("\t".join(row) for row in rows),
                method="native_spreadsheet",
                spreadsheet_preview=preview,
                rendered_html=csv_html,
            )

        # ── Helper: convert Excel date serial numbers to readable dates ──
        from datetime import datetime as _dt, timedelta as _td

        # Excel's epoch is 1899-12-30 (with the Lotus 1-2-3 1900 leap-year bug).
        # That makes 1900-01-01 a Sunday, so day-of-month directly indexes a
        # weekday name. Excel users frequently put weekdays in a column using a
        # custom format like `dddd` — openpyxl returns the underlying datetime,
        # not the rendered string, so we'd otherwise show "1900-01-01" etc. Map
        # 1900 dates back to weekday names — no real timesheet uses 1900 dates.
        _WEEKDAY_NAMES = ("Sunday", "Monday", "Tuesday", "Wednesday",
                          "Thursday", "Friday", "Saturday")

        def _format_cell(cell_value):
            """Convert a cell value to string, handling Excel date serials."""
            if cell_value is None:
                return ""
            if isinstance(cell_value, _dt):
                if cell_value.year == 1900:
                    # Day 1 = 1900-01-01 = Sunday; map to weekday name.
                    return _WEEKDAY_NAMES[(cell_value.day - 1) % 7]
                return cell_value.strftime("%Y-%m-%d")
            if isinstance(cell_value, (int, float)):
                # Excel date serial range: ~36526 (2000-01-01) to ~54789 (2050-01-01)
                # Avoid converting regular numbers like hours (0-24), IDs, etc.
                serial = float(cell_value)
                if 36526 <= serial <= 54789 and serial == int(serial):
                    try:
                        # Excel epoch: Jan 0, 1900 (with Lotus 1-2-3 leap year bug)
                        base = _dt(1899, 12, 30)
                        return (base + _td(days=int(serial))).strftime("%Y-%m-%d")
                    except (ValueError, OverflowError):
                        pass
            return str(cell_value)

        # Try openpyxl first (.xlsx), then fall back to xlrd (.xls)
        lines: list[str] = []
        sheets_preview: list[dict] = []
        try:
            import openpyxl
            workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            for sheet in workbook.worksheets:
                lines.append(f"=== Sheet: {sheet.title} ===")
                sheet_rows: list[list[str]] = []
                for row in sheet.iter_rows(values_only=True):
                    cells = [_format_cell(cell) for cell in row]
                    if any(cell.strip() for cell in cells):
                        lines.append("\t".join(cells))
                        sheet_rows.append(cells)
                sheets_preview.append({"name": sheet.title, "rows": sheet_rows})
        except Exception:
            # openpyxl failed — try xlrd for old .xls format
            try:
                import xlrd
                workbook = xlrd.open_workbook(file_contents=content)
                for sheet in workbook.sheets():
                    lines.append(f"=== Sheet: {sheet.name} ===")
                    sheet_rows: list[list[str]] = []
                    for row_idx in range(sheet.nrows):
                        cells = []
                        for col in range(sheet.ncols):
                            cell_type = sheet.cell_type(row_idx, col)
                            cell_value = sheet.cell_value(row_idx, col)
                            # xlrd cell type 3 = XL_CELL_DATE
                            if cell_type == 3:
                                try:
                                    date_tuple = xlrd.xldate_as_tuple(cell_value, workbook.datemode)
                                    if date_tuple[0] == 1900:
                                        cells.append(_WEEKDAY_NAMES[(date_tuple[2] - 1) % 7])
                                    elif date_tuple[0]:
                                        cells.append(_dt(*date_tuple[:3]).strftime("%Y-%m-%d"))
                                    else:
                                        cells.append(f"{date_tuple[3]:02d}:{date_tuple[4]:02d}")
                                except Exception:
                                    cells.append(_format_cell(cell_value))
                            else:
                                cells.append(_format_cell(cell_value))
                        if any(cell.strip() for cell in cells):
                            lines.append("\t".join(cells))
                            sheet_rows.append(cells)
                    sheets_preview.append({"name": sheet.name, "rows": sheet_rows})
            except ImportError:
                raise ValueError("Neither openpyxl nor xlrd available for spreadsheet extraction")

        if not lines:
            raise ValueError("Spreadsheet extraction produced no content")
        for sheet in sheets_preview:
            sheet["rows"] = _trim_trailing_empty_columns(sheet["rows"])
            sheet["blocks"] = [
                {"rows": block} for block in _split_into_blocks(sheet["rows"])
            ]
        preview = {"sheets": sheets_preview} if sheets_preview else None
        # HTML render is display-only — never let its failure break extraction.
        rendered_html = None
        try:
            if "openxmlformats" in mime_type or mime_type.endswith(".sheet"):
                from app.services.xlsx_render import render_xlsx_to_html
                rendered_html = render_xlsx_to_html(content)
            elif mime_type == "application/vnd.ms-excel":
                from app.services.xlsx_render import render_xls_to_html
                rendered_html = render_xls_to_html(content)
        except Exception as render_exc:
            logger.warning("xlsx HTML render failed (non-fatal): %s", render_exc)
        return ExtractionResult(
            text="\n".join(lines),
            method="native_spreadsheet",
            spreadsheet_preview=preview,
            rendered_html=rendered_html,
        )
    except Exception as exc:
        logger.warning("Spreadsheet extraction failed: %s", exc)
        return ExtractionResult(text="", method="failed", error=str(exc))


# ─── Word document extraction ────────────────────────────────────────────────

_DOCX_PREVIEW_CSS = """
:root { color-scheme: light dark; }
body {
    margin: 0;
    padding: 16px;
    font-family: 'Segoe UI', Calibri, Arial, sans-serif;
    font-size: 13px;
    background: transparent;
    color: #1a1a1a;
}
@media (prefers-color-scheme: dark) { body { color: #e6e6e6; } }
p { margin: 0 0 8px 0; }
table.docx-table {
    border-collapse: collapse;
    margin: 8px 0 16px 0;
}
table.docx-table td {
    border: 1px solid rgba(127,127,127,0.4);
    padding: 4px 8px;
    vertical-align: top;
}
""".strip()


def _wrap_docx_html(body: str) -> str:
    return (
        "<!doctype html><html><head><meta charset='utf-8'><style>"
        f"{_DOCX_PREVIEW_CSS}"
        f"</style></head><body>{body}</body></html>"
    )


def _dedupe_docx_row_cells(row) -> list:
    """
    Return one cell per visible column in a python-docx row, collapsing merges.

    python-docx exposes `row.cells` as the grid-expanded view: a horizontally
    merged cell appears once per underlying grid column, and each copy returns
    the full merged text. That's why a naive join produces
    "Total Weekly Hours:\tTotal Weekly Hours:\t..." for a row with a single
    merged header. Dedupe on the underlying `<w:tc>` XML element so each
    physical cell contributes text exactly once.
    """
    seen_ids: set[int] = set()
    unique: list = []
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen_ids:
            continue
        seen_ids.add(tc_id)
        unique.append(cell)
    return unique


def _extract_docx(content: bytes) -> ExtractionResult:
    """
    Extract text and an HTML preview from a .docx (Word 2007+) file.

    Plain text is what the downstream LLM sees — we join deduped row cells
    with tabs and paragraphs/rows with newlines. HTML is purely presentational
    for the review panel (mirrors what spreadsheet attachments produce via
    `rendered_html`) and falls back to `None` if anything goes wrong so a
    rendering bug can never mask a successful extraction.
    """
    try:
        import docx  # python-docx
        from html import escape as _html_escape

        document = docx.Document(io.BytesIO(content))
        text_parts: list[str] = []
        html_parts: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                html_parts.append(f"<p>{_html_escape(paragraph.text)}</p>")

        for table in document.tables:
            html_rows: list[str] = []
            for row in table.rows:
                unique_cells = _dedupe_docx_row_cells(row)
                cell_texts = [(cell.text or "").strip() for cell in unique_cells]
                if not any(cell_texts):
                    continue
                text_parts.append("\t".join(cell_texts))
                html_cells = "".join(
                    f"<td>{_html_escape(t).replace(chr(10), '<br/>')}</td>"
                    for t in cell_texts
                )
                html_rows.append(f"<tr>{html_cells}</tr>")
            if html_rows:
                html_parts.append(
                    '<table class="docx-table">' + "".join(html_rows) + "</table>"
                )

        text = "\n".join(text_parts).strip()
        if not text:
            return ExtractionResult(
                text="",
                method="failed",
                error="No text content found in .docx file.",
            )

        # Review panel renders this inside a sandboxed <iframe> (see
        # ReviewPanelPage.tsx), which doesn't inherit the parent's CSS. Emit a
        # full standalone HTML document with dark-mode-aware styling so the
        # preview is readable on both themes — same approach as xlsx_render.
        rendered_html = (
            _wrap_docx_html("".join(html_parts)) if html_parts else None
        )
        return ExtractionResult(
            text=text,
            method="native_spreadsheet",
            rendered_html=rendered_html,
        )
    except Exception as exc:
        logger.warning(".docx extraction failed: %s", exc)
        return ExtractionResult(text="", method="failed", error=str(exc))


async def _extract_doc(content: bytes) -> ExtractionResult:
    """
    Extract plain text from a legacy .doc (Word 97–2003 binary) file by
    shelling out to `antiword`. The subprocess is wrapped in a 60s timeout so
    a malformed file can't hang the worker — same defensive pattern we need
    anywhere that shells out to an external binary.
    """
    import asyncio
    import os
    import tempfile

    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
    except Exception as exc:
        return ExtractionResult(text="", method="failed", error=f".doc write failed: {exc}")

    try:
        proc = await asyncio.create_subprocess_exec(
            "antiword", tmp_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        try:
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=60.0)
        except asyncio.TimeoutError:
            proc.kill()
            await proc.wait()
            return ExtractionResult(
                text="",
                method="failed",
                error="antiword timed out after 60s processing .doc file.",
            )

        if proc.returncode != 0:
            err = (stderr or b"").decode("utf-8", errors="replace").strip()
            return ExtractionResult(
                text="",
                method="failed",
                error=f"antiword exited {proc.returncode}: {err or 'no stderr'}",
            )

        text = (stdout or b"").decode("utf-8", errors="replace").strip()
        if not text:
            return ExtractionResult(
                text="",
                method="failed",
                error="antiword produced no output for .doc file.",
            )
        return ExtractionResult(text=text, method="native_spreadsheet")
    except FileNotFoundError:
        return ExtractionResult(
            text="",
            method="failed",
            error="antiword binary is not installed — legacy .doc files cannot be extracted.",
        )
    except Exception as exc:
        logger.warning(".doc extraction failed: %s", exc)
        return ExtractionResult(text="", method="failed", error=str(exc))
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ─── Native PDF extraction ────────────────────────────────────────────────────

def _extract_native_pdf(content: bytes) -> ExtractionResult:
    try:
        import pdfplumber

        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)

        text = "\n".join(text_parts)
        if len(text.strip()) > 100:
            return ExtractionResult(text=text, method="native_pdf")
        return ExtractionResult(
            text="",
            method="failed",
            error="Insufficient embedded PDF text; likely scanned.",
        )
    except Exception as exc:
        logger.warning("Native PDF extraction failed: %s", exc)
        return ExtractionResult(text="", method="failed", error=str(exc))


# ─── PDF rasterization ────────────────────────────────────────────────────────

async def _rasterize_pdf(content: bytes) -> list[bytes]:
    """
    Rasterize PDF pages to PNG using pdftoppm (port of extraction_service.ts).
    Falls back to pdf2image if pdftoppm is not available.
    """
    import asyncio
    import os
    import tempfile

    MAX_PDF_PAGES = 50  # Safety limit

    # Quick page count check using pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            if len(pdf.pages) > MAX_PDF_PAGES:
                logger.warning("PDF has %d pages, capping rasterization at %d", len(pdf.pages), MAX_PDF_PAGES)
    except Exception:
        pass  # If we can't count pages, proceed with caution

    try:
        with tempfile.TemporaryDirectory() as tmp:
            input_path = os.path.join(tmp, "input.pdf")
            output_prefix = os.path.join(tmp, "page")

            with open(input_path, "wb") as f:
                f.write(content)

            proc = await asyncio.create_subprocess_exec(
                "pdftoppm", "-png", "-l", str(MAX_PDF_PAGES), input_path, output_prefix,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            await proc.communicate()

            if proc.returncode != 0:
                logger.warning("pdftoppm exited with code %d, falling back to pdf2image", proc.returncode)
                raise RuntimeError("pdftoppm failed")

            page_files = sorted(
                f for f in os.listdir(tmp)
                if f.startswith("page") and f.endswith(".png")
            )
            if page_files:
                result = []
                for p in page_files:
                    with open(os.path.join(tmp, p), "rb") as fh:
                        result.append(fh.read())
                return result
    except Exception:
        pass

    # Fallback to pdf2image
    from pdf2image import convert_from_bytes

    pages = convert_from_bytes(content, dpi=200, last_page=MAX_PDF_PAGES)
    result = []
    for page in pages:
        buf = io.BytesIO()
        page.save(buf, format="PNG")
        result.append(buf.getvalue())
    return result


# ─── Tesseract OCR ────────────────────────────────────────────────────────────

def _ocr_image(image_bytes: bytes, ocr_lang: str = "eng") -> tuple[str, float]:
    import pytesseract
    from PIL import Image

    image = Image.open(io.BytesIO(image_bytes))
    data = pytesseract.image_to_data(
        image,
        output_type=pytesseract.Output.DICT,
        lang=ocr_lang,
    )

    confidences: list[int] = []
    for raw_conf in data.get("conf", []):
        try:
            conf = int(float(raw_conf))
        except (TypeError, ValueError):
            continue
        if conf > 0:
            confidences.append(conf)

    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
    text = pytesseract.image_to_string(image, lang=ocr_lang)
    return text, avg_confidence


def _extract_tesseract_pdf(content: bytes, ocr_lang: str = "eng") -> tuple[str, float]:
    from pdf2image import convert_from_bytes

    MAX_PDF_PAGES = 50  # Safety limit
    pages = convert_from_bytes(content, dpi=200, last_page=MAX_PDF_PAGES)
    texts: list[str] = []
    confidences: list[float] = []

    for page_image in pages:
        buffer = io.BytesIO()
        page_image.save(buffer, format="PNG")
        text, confidence = _ocr_image(buffer.getvalue(), ocr_lang=ocr_lang)
        texts.append(text)
        confidences.append(confidence)

    average = sum(confidences) / len(confidences) if confidences else 0.0
    return "\n".join(texts), average


# ─── Vision API ───────────────────────────────────────────────────────────────

def _extract_vision_timesheets(payload: dict | list) -> list[dict]:
    """
    Port of extractVisionTimesheets from extraction_service.ts.
    Handles array, wrapped object, or single timesheet object.
    """
    if isinstance(payload, list):
        return [item for item in payload if item]
    if isinstance(payload, dict):
        if isinstance(payload.get("timesheets"), list):
            return [item for item in payload["timesheets"] if item]
        # Single timesheet object returned directly
        if (
            payload.get("employee_name") is not None
            or payload.get("period_start") is not None
            or payload.get("line_items") is not None
            or payload.get("total_hours") is not None
        ):
            return [payload]
    return []


async def _call_vision_api(
    image_bytes_list: list[bytes],
    mime_type: str = "image/png",
    reference_date: str | None = None,
) -> ExtractionResult:
    """
    Call the Vision API with one or more page images.
    Returns ExtractionResult with vision_timesheets populated (structured JSON).
    Port of extraction_service.ts runVisionFallback.
    """
    if not image_bytes_list:
        return ExtractionResult(text="", method="failed", error="No images to process.")

    try:
        from openai import AsyncOpenAI
    except ModuleNotFoundError as exc:
        logger.error("Vision extraction unavailable: %s", exc)
        return ExtractionResult(text="", method="failed", error=str(exc))

    if not settings.openai_api_key:
        return ExtractionResult(text="", method="failed", error="OPENAI_API_KEY is not configured.")

    try:
        client = AsyncOpenAI(api_key=settings.openai_api_key, timeout=60.0)

        content: list[dict] = []
        for image_bytes in image_bytes_list:
            encoded = base64.b64encode(image_bytes).decode("utf-8")
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{mime_type};base64,{encoded}"},
            })
        prompt = _build_vision_prompt(reference_date) if reference_date else VISION_PROMPT
        content.append({"type": "text", "text": prompt})

        response = await client.chat.completions.create(
            model="gpt-4o",
            max_tokens=min(4000 + len(image_bytes_list) * 500, 16000),
            messages=[{"role": "user", "content": content}],
        )

        raw = response.choices[0].message.content or ""
        try:
            # Strip markdown code fences if present
            clean = raw.strip()
            if clean.startswith("```"):
                clean = re.sub(r"^```[a-z]*\n?", "", clean)
                clean = re.sub(r"\n?```$", "", clean)
            parsed = json.loads(clean)
        except json.JSONDecodeError:
            logger.warning("Vision API returned non-JSON response, using as text: %s", raw[:200])
            return ExtractionResult(text=raw, method="vision_api")

        vision_timesheets: list[dict] = []
        if isinstance(parsed, dict):
            raw_timesheets = parsed.get("timesheets", [])
            if isinstance(raw_timesheets, list):
                validated = [ts for ts in raw_timesheets if _validate_vision_timesheet(ts)]
                if validated:
                    vision_timesheets = validated
        if not vision_timesheets:
            # Fall back to the existing flexible extractor for non-standard shapes
            vision_timesheets = [
                ts for ts in _extract_vision_timesheets(parsed)
                if _validate_vision_timesheet(ts)
            ]
        return ExtractionResult(
            text="",
            method="vision_api",
            vision_timesheets=vision_timesheets if vision_timesheets else None,
        )
    except Exception as exc:
        logger.error("Vision API extraction failed: %s", exc)
        return ExtractionResult(text="", method="failed", error=str(exc))


# Keep the import here so it's available after the function definition above
import base64  # noqa: E402


# ─── MIME type normalisation ──────────────────────────────────────────────────

def _normalize_mime_type(filename: str, mime_type: str) -> str:
    mime_lower = (mime_type or "").lower().strip()
    if mime_lower and mime_lower != "application/octet-stream":
        return mime_lower

    suffix = Path(filename).suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if suffix == ".csv":
        return "text/csv"
    if suffix == ".pdf":
        return "application/pdf"
    if suffix == ".png":
        return "image/png"
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix in {".tif", ".tiff"}:
        return "image/tiff"
    if suffix == ".bmp":
        return "image/bmp"
    if suffix == ".gif":
        return "image/gif"
    if suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix == ".doc":
        return "application/msword"
    return mime_lower


# ─── Main extraction entry point ─────────────────────────────────────────────

async def extract_text(
    content: bytes,
    filename: str,
    mime_type: str,
    reference_date: str | None = None,
) -> ExtractionResult:
    """
    Run the document extraction pipeline for a single attachment.

    reference_date: ISO-formatted date (e.g. '2026-04-17'), typically the email's
    received_at date. Passed to the Vision LLM so year-less dates in the document
    resolve to the correct year instead of defaulting to an old training-data year.

    Extraction order (port of extraction_service.ts extractFromBuffer):
      Spreadsheet  → native parser → text
      PDF          → native text → Vision (pdftoppm, all pages) → Tesseract fallback
      Image        → Vision → Tesseract fallback
    """
    normalized_mime = _normalize_mime_type(filename, mime_type)

    # ── Spreadsheet ───────────────────────────────────────────────────────────
    if normalized_mime in SPREADSHEET_MIME_TYPES:
        logger.debug("Extracting spreadsheet: %s", filename)
        return _extract_spreadsheet(content, normalized_mime)

    # ── Word (.docx) ──────────────────────────────────────────────────────────
    if normalized_mime in WORD_DOCX_MIME_TYPES:
        logger.debug("Extracting .docx: %s", filename)
        return _extract_docx(content)

    # ── Legacy Word (.doc) ────────────────────────────────────────────────────
    if normalized_mime in WORD_DOC_MIME_TYPES:
        logger.debug("Extracting .doc via antiword: %s", filename)
        return await _extract_doc(content)

    # ── PDF ───────────────────────────────────────────────────────────────────
    if "pdf" in normalized_mime:
        # 1. Native text extraction
        logger.debug("Trying native PDF extraction: %s", filename)
        result = _extract_native_pdf(content)
        if result.success:
            return result

        # 2. Rasterize with pdftoppm → Vision API (all pages)
        logger.debug("Native PDF extraction insufficient; trying Vision API: %s", filename)
        try:
            pages = await _rasterize_pdf(content)
            if pages:
                vision_result = await _call_vision_api(pages, reference_date=reference_date)
                if vision_result.success:
                    return vision_result
        except Exception as exc:
            logger.warning("PDF Vision extraction failed for %s: %s", filename, exc)

        # 3. Tesseract OCR fallback
        logger.debug("Vision failed; trying Tesseract OCR: %s", filename)
        existing_text = result.text if result else ""
        ocr_lang = _detect_tesseract_lang(existing_text) if existing_text and len(existing_text) > 20 else "eng"
        try:
            text, confidence = _extract_tesseract_pdf(content, ocr_lang=ocr_lang)
            if text.strip():
                return ExtractionResult(text=text, method="tesseract", confidence=confidence)
        except Exception as exc:
            logger.warning("PDF OCR failed for %s: %s", filename, exc)

        return ExtractionResult(text="", method="failed", error="All PDF extraction methods failed.")

    # ── Image ─────────────────────────────────────────────────────────────────
    if normalized_mime in IMAGE_MIME_TYPES:
        # 1. Vision API first (port of extraction_service.ts — Vision before Tesseract for images)
        logger.debug("Trying Vision API for image: %s", filename)
        try:
            vision_result = await _call_vision_api([content], mime_type=normalized_mime, reference_date=reference_date)
            if vision_result.success:
                return vision_result
        except Exception as exc:
            logger.warning("Vision API failed for %s: %s", filename, exc)

        # 2. Tesseract fallback
        logger.debug("Vision failed; trying Tesseract OCR for image: %s", filename)
        try:
            text, confidence = _ocr_image(content, ocr_lang="eng")
            if text.strip():
                return ExtractionResult(text=text, method="tesseract", confidence=confidence)
        except Exception as exc:
            logger.warning("Image OCR failed for %s: %s", filename, exc)

        return ExtractionResult(text="", method="failed", error="All image extraction methods failed.")

    return ExtractionResult(
        text="",
        method="failed",
        error=f"Unsupported MIME type: {mime_type or normalized_mime}",
    )
